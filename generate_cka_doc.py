#!/usr/bin/env python3
"""Generate a comprehensive CKA Exam Cheat Sheet Word Document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(9.5)

# Code style
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_font = code_style.font
code_font.name = 'Consolas'
code_font.size = Pt(8)
code_font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.space_after = Pt(2)
code_style.paragraph_format.left_indent = Cm(0.5)

# Heading styles
for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

doc.styles['Heading 1'].font.size = Pt(16)
doc.styles['Heading 1'].paragraph_format.space_before = Pt(14)
doc.styles['Heading 1'].paragraph_format.space_after = Pt(6)
doc.styles['Heading 2'].font.size = Pt(12)
doc.styles['Heading 2'].paragraph_format.space_before = Pt(10)
doc.styles['Heading 2'].paragraph_format.space_after = Pt(4)
doc.styles['Heading 3'].font.size = Pt(10)
doc.styles['Heading 3'].paragraph_format.space_before = Pt(8)
doc.styles['Heading 3'].paragraph_format.space_after = Pt(3)

def add_code(text, doc=doc):
    """Add a code block with gray background."""
    for line in text.strip().split('\n'):
        p = doc.add_paragraph(line, style='CodeBlock')
        # Add shading
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F0F0F0')
        shading.set(qn('w:val'), 'clear')
        p.paragraph_format.element.get_or_add_pPr().append(shading)

def add_yaml(text, doc=doc):
    """Add a YAML block."""
    add_code(text)

def add_tip(text, doc=doc):
    """Add a tip/note paragraph."""
    p = doc.add_paragraph()
    run = p.add_run('TIP: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    run.font.size = Pt(9)
    run2 = p.add_run(text)
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_exam_note(text, doc=doc):
    """Add an exam-specific note."""
    p = doc.add_paragraph()
    run = p.add_run('EXAM: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    run.font.size = Pt(9)
    run2 = p.add_run(text)
    run2.font.size = Pt(9)

def add_bullet(text, doc=doc):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style='List Bullet')
    p.style.font.size = Pt(9)

def add_section_break():
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(100)
run = p.add_run('CKA EXAM')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Comprehensive Cheat Sheet')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Certified Kubernetes Administrator')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run('All Imperative Commands | YAML Configs | Troubleshooting\n')
run.font.size = Pt(11)
run = p.add_run('Everything you need for the CKA exam & killer.sh practice')
run.font.size = Pt(11)
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('Exam: 2 hours | 66% to pass | One tab: kubernetes.io/docs')
run.font.size = Pt(12)
run.bold = True
run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
run = p.add_run('Exam Domains: Architecture 25% | Workloads 15% | Networking 20% | Storage 10% | Troubleshooting 30%')
run.font.size = Pt(10)

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Exam Environment Setup & Shell Aliases',
    '2. Imperative Commands — Speed Reference',
    '3. kubectl Quick Reference',
    '4. Core Concepts — Cluster Architecture',
    '5. Pod YAML Skeleton',
    '6. Deployment + Service YAML',
    '7. Scheduling — Taints, Tolerations, Affinity, nodeSelector',
    '8. Resource Limits, LimitRange, ResourceQuota',
    '9. DaemonSet, Job, CronJob',
    '10. Static Pods',
    '11. ConfigMaps & Secrets',
    '12. Multi-Container Pods & Init Containers',
    '13. Probes — Liveness & Readiness',
    '14. Rolling Updates & Rollbacks',
    '15. HPA — Autoscaling',
    '16. RBAC — Roles, Bindings, ServiceAccounts',
    '17. TLS & Certificates — CSR Workflow',
    '18. kubeconfig',
    '19. Security Contexts',
    '20. NetworkPolicy',
    '21. Storage — PV, PVC, StorageClass',
    '22. Services & Networking',
    '23. DNS & CoreDNS',
    '24. Ingress',
    '25. ETCD Backup & Restore',
    '26. Cluster Upgrade (kubeadm)',
    '27. Node Maintenance — Cordon, Drain, Uncordon',
    '28. PodDisruptionBudget',
    '29. Troubleshooting Checklist',
    '30. JSONPath & Output Formatting',
    '31. Helm',
    '32. Kustomize',
    '33. kubeadm Installation',
    '34. CRDs & Operators',
    '35. Practice Scenarios & Mock Questions',
    '36. Exam Tips & Common Mistakes',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 1. EXAM ENVIRONMENT SETUP
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('1. Exam Environment Setup & Shell Aliases', level=1)
doc.add_paragraph('Run these commands FIRST when the exam starts:')

doc.add_heading('Shell Aliases (~/.bashrc)', level=3)
add_code("""# Essential aliases — set these at the START of the exam
alias k=kubectl
alias kn='kubectl config set-context --current --namespace'
alias kgp='kubectl get pods'
alias kgs='kubectl get svc'
alias kgn='kubectl get nodes'
alias kga='kubectl get all'
alias kaf='kubectl apply -f'
alias kdp='kubectl describe pod'
alias kl='kubectl logs'
alias ke='kubectl exec -it'

# Enable kubectl autocompletion
source <(kubectl completion bash)
complete -o default -F __start_kubectl k

# Set default editor
export KUBE_EDITOR=vi""")

doc.add_heading('Vim Settings (~/.vimrc)', level=3)
add_code("""set tabstop=2
set shiftwidth=2
set expandtab
set number
set autoindent""")

add_exam_note('Switch context at the start of EVERY question: kubectl config use-context <context-name>')

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 2. IMPERATIVE COMMANDS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('2. Imperative Commands — Speed Reference', level=1)
doc.add_paragraph('Use imperative commands with --dry-run=client -o yaml to generate YAML fast, then edit as needed.')

doc.add_heading('Pod', level=3)
add_code("""k run <name> --image=<img>
k run <name> --image=<img> --restart=Never
k run <name> --image=<img> --dry-run=client -o yaml > pod.yaml
k run <name> --image=<img> --port=80
k run <name> --image=<img> --labels="app=web,env=prod"
k run <name> --image=<img> --command -- sleep 3600
k run <name> --image=<img> -- <arg1> <arg2>""")

doc.add_heading('Deployment', level=3)
add_code("""k create deployment <name> --image=<img> [--replicas=N]
k create deployment <name> --image=<img> --dry-run=client -o yaml > deploy.yaml
k scale deployment <name> --replicas=N
k set image deployment/<name> <container>=<img>
k rollout status deployment/<name>
k rollout history deployment/<name>
k rollout undo deployment/<name>
k rollout undo deployment/<name> --to-revision=2""")

doc.add_heading('Service', level=3)
add_code("""k expose deployment <name> --port=<port> [--name=<svc-name>]
k expose deployment <name> --port=80 --type=NodePort
k expose pod <name> --port=80
k create service nodeport nginx --tcp=80:80 --node-port=30080 --dry-run=client -o yaml
k expose deployment <name> --port=80 --type=ClusterIP --dry-run=client -o yaml""")

doc.add_heading('Namespace', level=3)
add_code("""k create namespace <name>
k create ns <name>""")

doc.add_heading('ConfigMap & Secret', level=3)
add_code("""k create configmap <name> --from-literal=KEY=val [--from-literal=KEY2=val2]
k create configmap <name> --from-file=path/to/file
k create configmap <name> --from-file=key=path
k create secret generic <name> --from-literal=KEY=val
k create secret generic <name> --from-file=key=path
k create secret tls <name> --cert=tls.crt --key=tls.key
k create secret docker-registry regcred \\
  --docker-server=<url> --docker-username=<user> \\
  --docker-password=<pass> --docker-email=<email>""")

doc.add_heading('ServiceAccount', level=3)
add_code("""k create serviceaccount <name>
k create sa <name>
k create token <sa-name>    # short-lived token (K8s 1.24+)""")

doc.add_heading('Job & CronJob', level=3)
add_code("""k create job <name> --image=busybox -- echo hello
k create job <name> --image=busybox --dry-run=client -o yaml -- echo hello
k create cronjob <name> --image=busybox --schedule="0 * * * *" -- echo hi
k create cronjob <name> --image=busybox --schedule="*/5 * * * *" --dry-run=client -o yaml -- echo hi""")

doc.add_heading('RBAC', level=3)
add_code("""k create role <name> --verb=get,list,create,delete --resource=pods -n <ns>
k create rolebinding <name> --role=<role> --user=<user> -n <ns>
k create rolebinding <name> --role=<role> --serviceaccount=<ns>:<sa> -n <ns>
k create clusterrole <name> --verb=get,list --resource=nodes
k create clusterrolebinding <name> --clusterrole=<name> --user=<user>
k create clusterrolebinding <name> --clusterrole=<name> --serviceaccount=<ns>:<sa>""")

doc.add_heading('Ingress', level=3)
add_code("""k create ingress <name> --rule="host/path=svc:80" --dry-run=client -o yaml""")

doc.add_heading('Generate → Edit → Apply Pattern', level=3)
add_code("""k run mypod --image=nginx --dry-run=client -o yaml > mypod.yaml
vi mypod.yaml   # add labels, resources, volumes, etc.
k apply -f mypod.yaml

# Edit running resource
k edit deployment nginx

# Force replace (immutable fields)
k replace --force -f pod.yaml

# Delete stuck pod
k delete pod x --force --grace-period=0""")

add_tip('Always add -n <namespace> or --namespace=<ns> for namespace-scoped resources.')

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 3. KUBECTL QUICK REFERENCE
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('3. kubectl Quick Reference', level=1)

doc.add_heading('Inspect & Debug', level=3)
add_code("""k get pods -A -o wide
k get all -n <ns>
k describe pod <name>
k logs <pod> [-c <container>] [--previous] [-f]
k exec -it <pod> -- sh
k exec -it <pod> -c <container> -- sh
k get events --sort-by='.lastTimestamp'
k get events -A --sort-by='.lastTimestamp'
k top nodes
k top pods --sort-by=memory
k top pods --sort-by=cpu""")

doc.add_heading('Docs in Terminal (Exam Allowed)', level=3)
add_code("""k explain pod
k explain pod.spec.containers
k explain pod.spec.containers.resources
k explain deployment.spec --recursive
k api-resources          # list all resource types
k api-resources --namespaced=true
k api-resources --namespaced=false
k api-versions           # list all API versions""")

doc.add_heading('Edit / Patch / Delete', level=3)
add_code("""k edit deploy <name>
k patch svc s1 -p '{"spec":{"type":"NodePort"}}'
k replace --force -f pod.yaml
k delete pod x --force --grace-period=0""")

doc.add_heading('Label / Annotate', level=3)
add_code("""k label pod <name> env=prod
k label pod <name> env=staging --overwrite
k label pod <name> env-                      # remove label
k label node <name> disk=ssd
k annotate pod <name> key=value
k get pods --selector app=web
k get pods --selector app=web,env=prod       # AND logic
k get pods --show-labels
k get pods -L app,env                        # show as columns
k get pods --selector env=production --no-headers | wc -l""")

doc.add_heading('Context & Namespace', level=3)
add_code("""k config use-context <ctx>
k config current-context
k config get-contexts
k config set-context --current --namespace=dev
k config view""")

doc.add_heading('Node Maintenance', level=3)
add_code("""k cordon <node>
k drain <node> --ignore-daemonsets --delete-emptydir-data
k drain <node> --ignore-daemonsets --delete-emptydir-data --force
k uncordon <node>""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 4. CORE CONCEPTS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('4. Core Concepts — Cluster Architecture', level=1)

doc.add_heading('Control Plane Components', level=2)
doc.add_paragraph('kube-apiserver: Front-end REST API. Only component that talks to etcd. Authenticates, authorizes, validates requests.')
doc.add_paragraph('etcd: Distributed key-value store. Uses RAFT consensus. Single source of truth. Backup with etcdctl snapshot.')
doc.add_paragraph('kube-scheduler: Watches unscheduled pods. Filtering (predicates) → Scoring (priorities) → Binding.')
doc.add_paragraph('kube-controller-manager: Runs controllers (Node, Replication, Endpoints). Reconciliation loop: current state → desired state.')
doc.add_paragraph('cloud-controller-manager: Cloud-specific logic (nodes, LBs, routes).')

doc.add_heading('Worker Node Components', level=2)
doc.add_paragraph('kubelet: Agent on each node. Registers node, ensures containers run. Reads PodSpecs from API server or static files.')
doc.add_paragraph('kube-proxy: Network proxy. Programs iptables/IPVS for Service routing.')
doc.add_paragraph('Container Runtime: containerd or CRI-O (Docker removed in K8s 1.24+). Uses CRI standard.')

doc.add_heading('crictl Commands (CRI debugging)', level=3)
add_code("""crictl ps -a                        # list all containers
crictl images                       # list images
crictl pods                         # list pods
crictl logs <container-id>          # view logs
crictl inspect <container-id>       # inspect container""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 5. POD YAML SKELETON
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('5. Pod YAML Skeleton (Complete)', level=1)
add_yaml("""apiVersion: v1
kind: Pod
metadata:
  name: mypod
  namespace: default
  labels:
    app: myapp
spec:
  serviceAccountName: my-sa
  nodeSelector:
    disk: ssd
  tolerations:
  - key: "node-role"
    operator: "Exists"
    effect: "NoSchedule"
  initContainers:
  - name: init
    image: busybox
    command: ['sh','-c','sleep 5']
  containers:
  - name: app
    image: nginx:1.21
    ports:
    - containerPort: 80
    command: ["sleep"]          # overrides ENTRYPOINT
    args: ["3600"]              # overrides CMD
    env:
    - name: KEY
      value: val
    - name: FROM_CM
      valueFrom:
        configMapKeyRef:
          name: my-cm
          key: APP_KEY
    - name: FROM_SECRET
      valueFrom:
        secretKeyRef:
          name: my-secret
          key: DB_PASS
    envFrom:
    - configMapRef:
        name: my-cm
    - secretRef:
        name: my-secret
    resources:
      requests: { cpu: 100m, memory: 128Mi }
      limits:   { cpu: 500m, memory: 256Mi }
    livenessProbe:
      httpGet: { path: /healthz, port: 80 }
      initialDelaySeconds: 10
      periodSeconds: 20
    readinessProbe:
      httpGet: { path: /ready, port: 80 }
      initialDelaySeconds: 5
      periodSeconds: 10
    volumeMounts:
    - name: data
      mountPath: /data
    - name: config-vol
      mountPath: /etc/config
    securityContext:
      runAsUser: 1000
      runAsNonRoot: true
      readOnlyRootFilesystem: true
      capabilities:
        add: ["NET_ADMIN"]
        drop: ["ALL"]
  volumes:
  - name: data
    persistentVolumeClaim:
      claimName: my-pvc
  - name: config-vol
    configMap:
      name: my-cm""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 6. DEPLOYMENT + SERVICE
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('6. Deployment + Service YAML', level=1)

doc.add_heading('Deployment', level=2)
add_yaml("""apiVersion: apps/v1
kind: Deployment
metadata:
  name: webapp
spec:
  replicas: 3
  strategy:
    type: RollingUpdate
    rollingUpdate: { maxSurge: 1, maxUnavailable: 0 }
  selector:
    matchLabels: { app: webapp }
  template:
    metadata:
      labels: { app: webapp }
    spec:
      containers:
      - name: webapp
        image: nginx:1.21
        ports: [{ containerPort: 80 }]
        resources:
          requests: { cpu: 100m, memory: 128Mi }
          limits:   { cpu: 500m, memory: 256Mi }""")

doc.add_heading('Service (NodePort)', level=2)
add_yaml("""apiVersion: v1
kind: Service
metadata:
  name: webapp-svc
spec:
  type: NodePort
  selector: { app: webapp }
  ports:
  - port: 80
    targetPort: 80
    nodePort: 30080""")

doc.add_heading('Service (ClusterIP)', level=2)
add_yaml("""apiVersion: v1
kind: Service
metadata:
  name: back-end
spec:
  type: ClusterIP
  ports:
  - targetPort: 80
    port: 80
  selector:
    app: myapp""")

doc.add_paragraph('Service Types: ClusterIP (internal, default) | NodePort (node port 30000-32767) | LoadBalancer (cloud LB)')

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 7. SCHEDULING
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('7. Scheduling — Taints, Tolerations, Affinity, nodeSelector', level=1)

doc.add_heading('Taints & Tolerations', level=2)
doc.add_paragraph('Taints on NODES repel pods. Tolerations on PODS allow them onto tainted nodes.')
doc.add_paragraph('Effects: NoSchedule | PreferNoSchedule | NoExecute (evicts existing pods)')

add_code("""# Apply a taint
k taint nodes node01 app=blue:NoSchedule

# Remove a taint (trailing minus)
k taint nodes node01 app=blue:NoSchedule-

# Remove control-plane taint (single-node clusters)
k taint nodes controlplane node-role.kubernetes.io/control-plane:NoSchedule-

# Check taints
k describe node node01 | grep -i taint""")

doc.add_heading('Toleration in Pod YAML', level=3)
add_yaml("""tolerations:
- key: "app"
  operator: "Equal"       # or "Exists" (matches any value)
  value: "blue"
  effect: "NoSchedule"
# NoExecute toleration with grace period:
- key: "node.kubernetes.io/not-ready"
  operator: "Exists"
  effect: "NoExecute"
  tolerationSeconds: 300""")

doc.add_heading('nodeSelector (Simplest)', level=2)
add_code("""# Label a node
k label nodes node01 size=Large
k label nodes node01 size-     # remove label""")
add_yaml("""spec:
  nodeSelector:
    size: Large""")

doc.add_heading('Node Affinity (Advanced)', level=2)
add_yaml("""spec:
  affinity:
    nodeAffinity:
      requiredDuringSchedulingIgnoredDuringExecution:   # HARD rule
        nodeSelectorTerms:
        - matchExpressions:
          - key: size
            operator: In          # In, NotIn, Exists, DoesNotExist, Gt, Lt
            values: [Large, Medium]
      preferredDuringSchedulingIgnoredDuringExecution:  # SOFT rule
      - weight: 50
        preference:
          matchExpressions:
          - key: disktype
            operator: In
            values: [ssd]""")

doc.add_heading('Manual Scheduling (nodeName)', level=2)
add_yaml("""spec:
  nodeName: node02        # bypasses scheduler entirely""")
add_exam_note('nodeName can only be set at creation time. To move a pod: export YAML → edit → delete → recreate.')

doc.add_heading('Dedicated Nodes Pattern (Taint + Affinity)', level=2)
add_code("""# Step 1: Taint the node
k taint nodes node01 dedicated=team-a:NoSchedule

# Step 2: Label the node
k label nodes node01 dedicated=team-a

# Step 3: Pod spec with BOTH toleration AND node affinity""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 8. RESOURCE LIMITS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('8. Resource Limits, LimitRange, ResourceQuota', level=1)

doc.add_heading('Resource Requests & Limits', level=2)
doc.add_paragraph('Requests = guaranteed allocation (scheduler uses). Limits = max at runtime.')
doc.add_paragraph('CPU > limit → throttled (never killed). Memory > limit → OOMKilled (exit 137).')
doc.add_paragraph('CPU: 1 = 1 vCPU, 500m = 0.5 CPU. Memory: Mi (mebibytes), Gi (gibibytes).')
add_yaml("""resources:
  requests:
    memory: "64Mi"
    cpu: "250m"
  limits:
    memory: "128Mi"
    cpu: "500m" """)

doc.add_paragraph('QoS Classes: Guaranteed (requests==limits) > Burstable (requests<limits) > BestEffort (no requests/limits)')

add_code("""k get pod x -o jsonpath='{.status.qosClass}'
k top nodes
k top pods --sort-by=memory
k describe node node01 | grep -A 10 "Allocated resources" """)

doc.add_heading('LimitRange (per-container defaults)', level=2)
add_yaml("""apiVersion: v1
kind: LimitRange
metadata:
  name: mem-cpu-limits
  namespace: dev
spec:
  limits:
  - type: Container
    default:          { cpu: "500m", memory: "128Mi" }
    defaultRequest:   { cpu: "250m", memory: "64Mi" }
    max:              { cpu: "2", memory: "1Gi" }
    min:              { cpu: "100m", memory: "16Mi" }""")

doc.add_heading('ResourceQuota (namespace-total)', level=2)
add_yaml("""apiVersion: v1
kind: ResourceQuota
metadata:
  name: compute-quota
  namespace: dev
spec:
  hard:
    pods: "10"
    requests.cpu: "4"
    requests.memory: 4Gi
    limits.cpu: "8"
    limits.memory: 8Gi
    services: "5"
    persistentvolumeclaims: "4" """)

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 9. DAEMONSET, JOB, CRONJOB
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('9. DaemonSet, Job, CronJob', level=1)

doc.add_heading('DaemonSet', level=2)
doc.add_paragraph('One pod per node (all or subset via nodeSelector). Use for: log collectors, node monitoring, CNI/CSI plugins.')
add_yaml("""apiVersion: apps/v1
kind: DaemonSet
metadata:
  name: monitoring-daemon
  namespace: kube-system
spec:
  selector:
    matchLabels: { app: monitoring-daemon }
  template:
    metadata:
      labels: { app: monitoring-daemon }
    spec:
      tolerations:
      - key: node-role.kubernetes.io/control-plane
        operator: Exists
        effect: NoSchedule
      containers:
      - name: agent
        image: prom/node-exporter
        resources:
          limits: { memory: "200Mi", cpu: "100m" }""")

add_code("""k get ds -A
k describe ds monitoring-daemon -n kube-system
k rollout undo ds monitoring-daemon -n kube-system""")

doc.add_heading('Job', level=2)
add_yaml("""apiVersion: batch/v1
kind: Job
metadata: { name: pi-job }
spec:
  completions: 3
  parallelism: 2
  backoffLimit: 4
  template:
    spec:
      containers:
      - name: pi
        image: perl
        command: ["perl","-Mbignum=bpi","-wle","print bpi(2000)"]
      restartPolicy: Never""")

doc.add_heading('CronJob', level=2)
add_yaml("""apiVersion: batch/v1
kind: CronJob
metadata: { name: backup }
spec:
  schedule: "0 2 * * *"
  successfulJobsHistoryLimit: 3
  failedJobsHistoryLimit: 1
  jobTemplate:
    spec:
      template:
        spec:
          containers:
          - name: backup
            image: busybox
            command: ["/bin/sh","-c","date"]
          restartPolicy: OnFailure""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 10. STATIC PODS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('10. Static Pods', level=1)
doc.add_paragraph('Managed by kubelet directly from /etc/kubernetes/manifests/. Not through API server. Control-plane components are static pods via kubeadm.')
doc.add_paragraph('To delete: remove the manifest file. kubectl delete will NOT work permanently.')
doc.add_paragraph('Naming: node name appended (e.g., kube-apiserver-controlplane).')

add_code("""# Find static pod path
ps aux | grep kubelet | grep config
cat /var/lib/kubelet/config.yaml | grep staticPodPath
# Default: /etc/kubernetes/manifests/

# Create a static pod
cat <<EOF > /etc/kubernetes/manifests/static-nginx.yaml
apiVersion: v1
kind: Pod
metadata:
  name: static-nginx
spec:
  containers:
  - name: nginx
    image: nginx:1.25
    ports:
    - containerPort: 80
EOF

# Delete a static pod
rm /etc/kubernetes/manifests/static-nginx.yaml

# Verify it's static (ownerReferences.kind == Node)
k get pod kube-apiserver-controlplane -n kube-system \\
  -o jsonpath='{.metadata.ownerReferences[*].kind}'""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 11. CONFIGMAPS & SECRETS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('11. ConfigMaps & Secrets', level=1)

doc.add_heading('ConfigMap', level=2)
add_code("""k create configmap app-config --from-literal=APP_COLOR=blue --from-literal=APP_MODE=prod
k create configmap app-config --from-file=app.properties
k get cm
k describe cm app-config""")

add_yaml("""apiVersion: v1
kind: ConfigMap
metadata:
  name: app-config
data:
  APP_COLOR: blue
  APP_MODE: prod""")

doc.add_heading('Inject into Pod', level=3)
add_yaml("""# All keys as env vars
envFrom:
- configMapRef:
    name: app-config

# Single key
env:
- name: APP_COLOR
  valueFrom:
    configMapKeyRef:
      name: app-config
      key: APP_COLOR

# As volume (files)
volumes:
- name: config-vol
  configMap:
    name: app-config""")

doc.add_heading('Secret', level=2)
add_code("""k create secret generic app-secret --from-literal=DB_Host=mysql --from-literal=DB_Password=pass123
k create secret tls webapp-tls --cert=tls.crt --key=tls.key
k create secret docker-registry regcred \\
  --docker-server=private-registry.io \\
  --docker-username=user --docker-password=pass --docker-email=user@org.com
k get secret app-secret -o yaml
echo -n "bX1zcWw=" | base64 --decode     # decode secret value
echo -n "myvalue" | base64               # encode for YAML""")

add_yaml("""apiVersion: v1
kind: Secret
metadata:
  name: app-secret
data:
  DB_Host: bX1zcWw=          # base64 encoded
  DB_Password: cGFzd3Jk

# Inject same as ConfigMap:
# envFrom: - secretRef: name: app-secret
# OR env: - valueFrom: secretKeyRef: name/key""")

doc.add_heading('Pull from Private Registry', level=3)
add_yaml("""spec:
  containers:
  - name: app
    image: private-registry.io/apps/internal-app
  imagePullSecrets:
  - name: regcred""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 12. MULTI-CONTAINER & INIT CONTAINERS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('12. Multi-Container Pods & Init Containers', level=1)

doc.add_heading('Patterns', level=2)
doc.add_paragraph('Sidecar: auxiliary (log shipper, proxy). Ambassador: proxy to external. Adapter: transforms output.')

doc.add_heading('Sidecar Example', level=3)
add_yaml("""apiVersion: v1
kind: Pod
metadata:
  name: app-with-sidecar
spec:
  containers:
  - name: app
    image: busybox
    command: ['sh', '-c', 'while true; do echo "$(date) INFO running" >> /var/log/app.log; sleep 5; done']
    volumeMounts:
    - name: log-volume
      mountPath: /var/log
  - name: sidecar
    image: busybox
    command: ['sh', '-c', 'tail -f /var/log/app.log']
    volumeMounts:
    - name: log-volume
      mountPath: /var/log
  volumes:
  - name: log-volume
    emptyDir: {}""")

doc.add_heading('Init Containers', level=2)
doc.add_paragraph('Run to completion BEFORE main containers. Sequential. Failure → pod restarts.')
add_yaml("""spec:
  initContainers:
  - name: init-myservice
    image: busybox:1.28
    command: ['sh', '-c', 'until nslookup myservice; do echo waiting; sleep 2; done']
  - name: init-mydb
    image: busybox:1.28
    command: ['sh', '-c', 'until nslookup mydb; do echo waiting; sleep 2; done']
  containers:
  - name: myapp
    image: busybox:1.28
    command: ['sh', '-c', 'echo running && sleep 3600']""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 13. PROBES
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('13. Probes — Liveness & Readiness', level=1)
doc.add_paragraph('livenessProbe: fails → container restarted. readinessProbe: fails → removed from Service endpoints.')
doc.add_paragraph('Types: httpGet, exec, tcpSocket')

add_yaml("""containers:
- name: app
  image: myapp:1.0
  livenessProbe:
    httpGet:
      path: /healthz
      port: 8080
    initialDelaySeconds: 15
    periodSeconds: 20
  readinessProbe:
    httpGet:
      path: /ready
      port: 8080
    initialDelaySeconds: 5
    periodSeconds: 10

# Exec probe:
  livenessProbe:
    exec:
      command: ["cat", "/tmp/healthy"]
    initialDelaySeconds: 5

# TCP probe:
  readinessProbe:
    tcpSocket:
      port: 3306
    initialDelaySeconds: 10""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 14. ROLLING UPDATES & ROLLBACKS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('14. Rolling Updates & Rollbacks', level=1)
doc.add_paragraph('Strategies: RollingUpdate (default, gradual) | Recreate (kill all, then create new)')

add_code("""k rollout status deployment/myapp
k rollout history deployment/myapp
k set image deployment/myapp nginx=nginx:1.25
k rollout undo deployment/myapp
k rollout undo deployment/myapp --to-revision=2
k rollout pause deployment/myapp
k rollout resume deployment/myapp""")

add_yaml("""spec:
  strategy:
    type: RollingUpdate
    rollingUpdate:
      maxSurge: 1           # max extra pods during rollout
      maxUnavailable: 0     # all existing pods stay until new ones ready""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 15. HPA
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('15. HPA — Horizontal Pod Autoscaler', level=1)
doc.add_paragraph('Scales Deployment replicas based on CPU/memory. Requires Metrics Server. Pods MUST have resources.requests.')

add_code("""k autoscale deployment myapp --min=2 --max=10 --cpu-percent=80
k get hpa
k describe hpa myapp""")

add_yaml("""apiVersion: autoscaling/v2
kind: HorizontalPodAutoscaler
metadata:
  name: myapp-hpa
spec:
  scaleTargetRef:
    apiVersion: apps/v1
    kind: Deployment
    name: myapp
  minReplicas: 2
  maxReplicas: 10
  metrics:
  - type: Resource
    resource:
      name: cpu
      target:
        type: Utilization
        averageUtilization: 80""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 16. RBAC
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('16. RBAC — Roles, Bindings, ServiceAccounts', level=1)

doc.add_heading('Imperative (Fastest in Exam)', level=2)
add_code("""k create role dev-role --verb=get,list,create,delete --resource=pods -n dev
k create rolebinding dev-bind --role=dev-role --user=jane -n dev

k create clusterrole node-viewer --verb=get,list --resource=nodes
k create clusterrolebinding node-bind --clusterrole=node-viewer --user=jane

# Check permissions
k auth can-i create pods --as=jane -n dev
k auth can-i '*' '*' --as=system:serviceaccount:default:my-sa
k auth can-i list pods -n dev --as system:serviceaccount:dev:app-sa""")

doc.add_heading('Role YAML', level=2)
add_yaml("""apiVersion: rbac.authorization.k8s.io/v1
kind: Role
metadata:
  name: developer
  namespace: dev
rules:
- apiGroups: [""]
  resources: ["pods"]
  verbs: ["get", "list", "create", "delete"]
- apiGroups: [""]
  resources: ["configmaps"]
  verbs: ["create"]""")

doc.add_heading('RoleBinding YAML', level=2)
add_yaml("""apiVersion: rbac.authorization.k8s.io/v1
kind: RoleBinding
metadata:
  name: dev-binding
  namespace: dev
subjects:
- kind: User
  name: jane
  apiGroup: rbac.authorization.k8s.io
roleRef:
  kind: Role
  name: developer
  apiGroup: rbac.authorization.k8s.io""")

doc.add_heading('ClusterRole + ClusterRoleBinding', level=2)
add_yaml("""apiVersion: rbac.authorization.k8s.io/v1
kind: ClusterRole
metadata:
  name: cluster-admin-role
rules:
- apiGroups: [""]
  resources: ["nodes"]
  verbs: ["get", "list", "delete", "create"]
---
apiVersion: rbac.authorization.k8s.io/v1
kind: ClusterRoleBinding
metadata:
  name: cluster-admin-binding
subjects:
- kind: User
  name: cluster-admin
  apiGroup: rbac.authorization.k8s.io
roleRef:
  kind: ClusterRole
  name: cluster-admin-role
  apiGroup: rbac.authorization.k8s.io""")

doc.add_heading('ServiceAccount', level=2)
add_code("""k create sa dashboard-sa
k create token dashboard-sa    # short-lived token (1.24+)""")
add_yaml("""spec:
  serviceAccountName: dashboard-sa
  automountServiceAccountToken: false   # opt-out""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 17. TLS & CERTIFICATES
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('17. TLS & Certificates — CSR Workflow', level=1)

doc.add_heading('Key Certificate Files (kubeadm)', level=2)
doc.add_paragraph('All under /etc/kubernetes/pki/:')
add_bullet('ca.crt / ca.key — Cluster CA (root of trust)')
add_bullet('apiserver.crt / apiserver.key — API server TLS')
add_bullet('apiserver-kubelet-client.crt — API server → kubelet client cert')
add_bullet('apiserver-etcd-client.crt — API server → etcd client cert')
add_bullet('etcd/ca.crt, etcd/server.crt — etcd own CA and server cert')

doc.add_heading('Certificate Inspection', level=2)
add_code("""openssl x509 -in /etc/kubernetes/pki/apiserver.crt -text -noout
# Check: Not Before / Not After (expiry), Subject, Issuer, SAN
openssl x509 -in /etc/kubernetes/pki/apiserver.crt -text -noout | grep -A2 Validity""")

doc.add_heading('Create User Certificate (CSR Workflow)', level=2)
add_code("""# 1. Generate key and CSR
openssl genrsa -out jane.key 2048
openssl req -new -key jane.key -subj "/CN=jane/O=dev" -out jane.csr

# 2. Base64 encode the CSR
cat jane.csr | base64 | tr -d '\\n'""")

add_yaml("""# 3. CSR Object
apiVersion: certificates.k8s.io/v1
kind: CertificateSigningRequest
metadata: { name: jane }
spec:
  request: <base64-encoded-csr>
  signerName: kubernetes.io/kube-apiserver-client
  usages: [client auth]""")

add_code("""# 4. Approve and extract
k certificate approve jane
k get csr jane -o jsonpath='{.status.certificate}' | base64 -d > jane.crt

# 5. Configure kubeconfig
k config set-credentials jane --client-certificate=jane.crt --client-key=jane.key
k config set-context jane-ctx --cluster=kubernetes --user=jane
k config use-context jane-ctx

# Deny a CSR
k certificate deny bad-user""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 18. KUBECONFIG
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('18. kubeconfig', level=1)
doc.add_paragraph('Default: ~/.kube/config. Contains: clusters (API server + CA), users (certs/tokens), contexts (cluster + user + ns).')

add_code("""k config view
k config view --kubeconfig=my-config
k config get-contexts
k config current-context
k config use-context <context>
k config set-context --current --namespace=dev
k config set-credentials user --client-certificate=user.crt --client-key=user.key
k config set-context user-ctx --cluster=kubernetes --user=user""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 19. SECURITY CONTEXTS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('19. Security Contexts', level=1)
doc.add_paragraph('Pod-level or container-level. Container overrides pod. Capabilities only at container level.')

add_yaml("""spec:
  securityContext:           # Pod-level
    runAsUser: 1000
    runAsGroup: 3000
    fsGroup: 2000
  containers:
  - name: ubuntu
    image: ubuntu
    securityContext:          # Container-level (overrides pod)
      runAsUser: 1000
      runAsNonRoot: true
      readOnlyRootFilesystem: true
      capabilities:
        add: ["NET_ADMIN", "SYS_TIME"]
        drop: ["ALL"]""")

add_code("""k exec my-pod -- whoami
k exec my-pod -- id""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 20. NETWORK POLICY
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('20. NetworkPolicy', level=1)
doc.add_paragraph('Controls ingress/egress to pods. Requires CNI with policy support (Calico, Cilium). Default: allow all. Once policy selects pod → deny all not explicitly allowed.')

doc.add_heading('Default Deny All', level=2)
add_yaml("""apiVersion: networking.k8s.io/v1
kind: NetworkPolicy
metadata:
  name: default-deny-all
  namespace: prod
spec:
  podSelector: {}
  policyTypes:
  - Ingress
  - Egress""")

doc.add_heading('Allow Specific Ingress', level=2)
add_yaml("""apiVersion: networking.k8s.io/v1
kind: NetworkPolicy
metadata:
  name: allow-api-to-db
  namespace: prod
spec:
  podSelector:
    matchLabels: { app: db }
  policyTypes: [Ingress, Egress]
  ingress:
  - from:
    - podSelector:
        matchLabels: { app: api }
    - namespaceSelector:
        matchLabels: { env: prod }
    ports: [{ protocol: TCP, port: 3306 }]
  egress:
  - to:
    - podSelector:
        matchLabels: { app: cache }
    ports: [{ protocol: TCP, port: 6379 }]""")

doc.add_heading('Cross-namespace & IP Block', level=3)
add_yaml("""# Cross-namespace access
ingress:
- from:
  - namespaceSelector:
      matchLabels:
        name: monitoring

# External CIDR access
ingress:
- from:
  - ipBlock:
      cidr: 203.0.113.0/24
      except:
      - 203.0.113.128/25""")

add_code("""# Test from debug pod
k run test --rm -it --image=busybox:1.28 --restart=Never -- sh
nc -zv db 3306
wget -qO- http://web.default.svc.cluster.local:80""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 21. STORAGE
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('21. Storage — PV, PVC, StorageClass', level=1)

doc.add_heading('PersistentVolume (PV)', level=2)
add_yaml("""apiVersion: v1
kind: PersistentVolume
metadata:
  name: pv1
spec:
  capacity: { storage: 1Gi }
  accessModes: [ReadWriteOnce]        # RWO, ROX, RWX
  persistentVolumeReclaimPolicy: Retain   # Retain, Delete, Recycle
  hostPath: { path: /mnt/data }""")

doc.add_heading('PersistentVolumeClaim (PVC)', level=2)
add_yaml("""apiVersion: v1
kind: PersistentVolumeClaim
metadata:
  name: pvc1
spec:
  accessModes: [ReadWriteOnce]
  resources:
    requests: { storage: 500Mi }
  storageClassName: ""     # empty = static binding only""")

doc.add_heading('Using PVC in Pod', level=2)
add_yaml("""spec:
  containers:
  - name: app
    image: nginx
    volumeMounts:
    - name: data
      mountPath: /data
  volumes:
  - name: data
    persistentVolumeClaim:
      claimName: pvc1""")

doc.add_heading('StorageClass (Dynamic Provisioning)', level=2)
add_yaml("""apiVersion: storage.k8s.io/v1
kind: StorageClass
metadata: { name: fast }
provisioner: kubernetes.io/gce-pd
parameters: { type: pd-ssd }
reclaimPolicy: Delete
allowVolumeExpansion: true""")

add_code("""k get pv
k get pvc
k get sc
k describe pv pv1
k describe pvc pvc1""")

doc.add_paragraph('AccessModes: RWO (ReadWriteOnce) | ROX (ReadOnlyMany) | RWX (ReadWriteMany)')
doc.add_paragraph('Volume Types: emptyDir (temp), hostPath (node), persistentVolumeClaim (persist), configMap, secret')

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 22. SERVICES & NETWORKING
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('22. Services & Networking', level=1)

doc.add_paragraph('Every Pod gets unique IP. Pods talk without NAT. Service = stable VIP + load balancing.')

doc.add_heading('Service Types', level=2)
doc.add_paragraph('ClusterIP (default): internal VIP. NodePort: node IP + port (30000-32767). LoadBalancer: cloud LB.')

doc.add_heading('Linux Networking Commands (CKA)', level=2)
add_code("""ip link                              # list interfaces
ip addr                              # show IPs
ip route show                        # routing table
ip route add 192.168.2.0/24 via 192.168.1.1
ip route add default via 192.168.1.1
cat /proc/sys/net/ipv4/ip_forward    # 0=disabled, 1=enabled
echo 1 > /proc/sys/net/ipv4/ip_forward""")

doc.add_heading('CNI (Container Network Interface)', level=2)
add_code("""ls /etc/cni/net.d/                   # CNI config
ls /opt/cni/bin/                     # CNI binaries
k get node -o jsonpath='{.spec.podCIDR}'""")

doc.add_heading('Service Debugging', level=2)
add_code("""k get svc,endpoints <svc-name>
# Empty endpoints → wrong selector
iptables-save | grep <service-name>
k logs -n kube-system -l k8s-app=kube-proxy""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 23. DNS & COREDNS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('23. DNS & CoreDNS', level=1)

doc.add_paragraph('CoreDNS: cluster internal DNS. Runs as pods in kube-system. Config in coredns ConfigMap (Corefile).')
doc.add_paragraph('FQDN: Service → my-svc.my-ns.svc.cluster.local | Pod → 1-2-3-4.my-ns.pod.cluster.local')

add_code("""# Test DNS
k run -it --rm debug --image=busybox:1.28 -- nslookup kubernetes
k run -it --rm debug --image=busybox:1.28 -- nslookup <svc>.<ns>.svc.cluster.local

# Check CoreDNS
k get pods -n kube-system -l k8s-app=kube-dns
k get configmap coredns -n kube-system -o yaml
k logs <coredns-pod> -n kube-system""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 24. INGRESS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('24. Ingress', level=1)
doc.add_paragraph('Requires Ingress Controller (e.g., nginx-ingress). Manages external HTTP/HTTPS access.')

doc.add_heading('Path-based Routing', level=2)
add_yaml("""apiVersion: networking.k8s.io/v1
kind: Ingress
metadata:
  name: app-ingress
  annotations:
    nginx.ingress.kubernetes.io/rewrite-target: /
spec:
  rules:
  - http:
      paths:
      - path: /wear
        pathType: Prefix
        backend:
          service: { name: wear-service, port: { number: 80 } }
      - path: /watch
        pathType: Prefix
        backend:
          service: { name: watch-service, port: { number: 80 } }""")

doc.add_heading('Host-based Routing', level=2)
add_yaml("""spec:
  rules:
  - host: wear.example.com
    http:
      paths:
      - path: /
        pathType: Prefix
        backend:
          service: { name: wear-svc, port: { number: 80 } }
  - host: watch.example.com
    http:
      paths:
      - path: /
        pathType: Prefix
        backend:
          service: { name: watch-svc, port: { number: 80 } }""")

doc.add_heading('Ingress with TLS', level=2)
add_code("""k create secret tls webapp-tls --cert=tls.crt --key=tls.key -n apps""")
add_yaml("""spec:
  tls:
  - hosts:
    - webapp.example.com
    secretName: webapp-tls
  rules:
  - host: webapp.example.com
    http:
      paths:
      - path: /
        pathType: Prefix
        backend:
          service: { name: webapp-svc, port: { number: 80 } }""")

add_code("""k create ingress my-ingress --rule="host/path=svc:80" --dry-run=client -o yaml
k get ingress
k describe ingress app-ingress""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 25. ETCD BACKUP & RESTORE
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('25. ETCD Backup & Restore', level=1)
add_exam_note('This appears in almost every CKA exam. Memorize the command with cert paths.')

doc.add_heading('Backup', level=2)
add_code("""ETCDCTL_API=3 etcdctl snapshot save /tmp/etcd.db \\
  --endpoints=https://127.0.0.1:2379 \\
  --cacert=/etc/kubernetes/pki/etcd/ca.crt \\
  --cert=/etc/kubernetes/pki/etcd/server.crt \\
  --key=/etc/kubernetes/pki/etcd/server.key

# Verify
ETCDCTL_API=3 etcdctl snapshot status /tmp/etcd.db""")

doc.add_heading('Restore', level=2)
add_code("""# Restore to new data directory
etcdutl snapshot restore /tmp/etcd.db --data-dir=/var/lib/etcd-restored

# Update etcd static pod manifest
vi /etc/kubernetes/manifests/etcd.yaml
# Change: hostPath.path → /var/lib/etcd-restored

# Restart kubelet
systemctl daemon-reload
systemctl restart kubelet""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 26. CLUSTER UPGRADE
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('26. Cluster Upgrade (kubeadm)', level=1)
doc.add_paragraph('Rule: Upgrade one minor version at a time. Control plane first, then workers.')

doc.add_heading('Control Plane', level=2)
add_code("""# 1. Upgrade kubeadm
apt-mark unhold kubeadm
apt-get update && apt-get install -y kubeadm=1.XX.0-*
apt-mark hold kubeadm

# 2. Plan and apply
kubeadm upgrade plan
kubeadm upgrade apply v1.XX.0

# 3. Drain control plane node
kubectl drain <cp-node> --ignore-daemonsets --delete-emptydir-data

# 4. Upgrade kubelet and kubectl
apt-mark unhold kubelet kubectl
apt-get install -y kubelet=1.XX.0-* kubectl=1.XX.0-*
apt-mark hold kubelet kubectl
systemctl daemon-reload && systemctl restart kubelet

# 5. Uncordon
kubectl uncordon <cp-node>""")

doc.add_heading('Worker Nodes', level=2)
add_code("""# From control plane:
kubectl drain <worker> --ignore-daemonsets --delete-emptydir-data

# On worker node:
apt-mark unhold kubeadm kubelet
apt-get install -y kubeadm=1.XX.0-* kubelet=1.XX.0-*
kubeadm upgrade node
systemctl daemon-reload && systemctl restart kubelet
apt-mark hold kubeadm kubelet

# From control plane:
kubectl uncordon <worker>""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 27. NODE MAINTENANCE
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('27. Node Maintenance — Cordon, Drain, Uncordon', level=1)

doc.add_paragraph('cordon = no new pods | drain = cordon + evict | uncordon = allow scheduling again')
doc.add_paragraph('If node is down > 5 min, pods may be terminated by controller.')

add_code("""kubectl cordon <node>                          # mark unschedulable
kubectl drain <node> --ignore-daemonsets --delete-emptydir-data
kubectl drain <node> --ignore-daemonsets --delete-emptydir-data --force
kubectl uncordon <node>                        # re-enable scheduling

# Typical workflow:
# 1. drain → 2. perform maintenance → 3. uncordon → 4. verify with kubectl get nodes""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 28. PDB
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('28. PodDisruptionBudget (PDB)', level=1)
doc.add_paragraph('Limits voluntary disruptions. drain respects PDBs.')

add_yaml("""apiVersion: policy/v1
kind: PodDisruptionBudget
metadata:
  name: myapp-pdb
spec:
  minAvailable: 2       # OR maxUnavailable: 1
  selector:
    matchLabels: { app: myapp }""")

add_code("""k get pdb
k describe pdb myapp-pdb""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 29. TROUBLESHOOTING
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('29. Troubleshooting Checklist (30% of Exam)', level=1)

doc.add_heading('Application Failure', level=2)
add_code("""# 1. Check pod status
k get pods -o wide
k describe pod <name>

# 2. Check logs (current + previous crashed instance)
k logs <pod> -c <container> --previous

# 3. Check Service & Endpoints
k get svc,endpoints <svc-name>
# Empty endpoints → wrong selector. Compare:
k describe svc <name>      # see Selector
k get pods --show-labels    # see pod labels

# 4. Exec into pod
k exec -it <pod> -- sh""")

doc.add_heading('Control Plane Failure', level=2)
add_code("""# Static pods in kube-system
k get pods -n kube-system
k logs kube-apiserver-master -n kube-system
k logs kube-scheduler-master -n kube-system
k logs kube-controller-manager-master -n kube-system

# Check etcd
ETCDCTL_API=3 etcdctl endpoint health \\
  --endpoints=https://127.0.0.1:2379 \\
  --cacert=/etc/kubernetes/pki/etcd/ca.crt \\
  --cert=/etc/kubernetes/pki/etcd/server.crt \\
  --key=/etc/kubernetes/pki/etcd/server.key""")

doc.add_heading('Worker Node Failure', level=2)
add_code("""k get nodes
k describe node <node>        # check Conditions, LastHeartbeatTime

# On the node:
systemctl status kubelet
journalctl -u kubelet -f
systemctl status containerd

# Common fixes:
systemctl start kubelet
systemctl daemon-reload && systemctl restart kubelet

# Check kubelet config
cat /var/lib/kubelet/config.yaml
cat /etc/kubernetes/kubelet.conf

# Check certs
openssl x509 -in /var/lib/kubelet/<node>.crt -text -noout""")

doc.add_heading('Network Troubleshooting', level=2)
add_code("""# 1. Check CNI pods
k get pods -n kube-system   # Calico/Flannel/Weave running?

# 2. Check CoreDNS
k get pods -n kube-system -l k8s-app=kube-dns

# 3. Test DNS
k run dns-test --image=busybox:1.28 --rm -it --restart=Never -- nslookup kubernetes

# 4. Check kube-proxy
k get ds kube-proxy -n kube-system
iptables -L -t nat | grep <svc>

# 5. Check IP forwarding
cat /proc/sys/net/ipv4/ip_forward

# 6. Debug from inside cluster
k run netshoot --image=nicolaka/netshoot -it --rm --restart=Never -- bash

# 7. Required ports: 6443 (API), 2379-2380 (etcd), 10250-10252 (kubelet,etc), 30000-32767 (NodePort)

# 8. Certificate validity
openssl x509 -in /etc/kubernetes/pki/apiserver.crt -text -noout | grep -A2 Validity""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 30. JSONPATH & OUTPUT
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('30. JSONPath & Output Formatting', level=1)

add_code("""# Single value
k get pod x -o jsonpath='{.spec.nodeName}'

# All pod names
k get pods -o jsonpath='{.items[*].metadata.name}'

# Range (formatted output)
k get nodes -o jsonpath='{range .items[*]}{.metadata.name}{"\\t"}{.status.capacity.cpu}{"\\n"}{end}'

# Filter
k get nodes -o jsonpath='{.items[*].status.addresses[?(@.type=="InternalIP")].address}'

# Custom columns
k get pods -A -o custom-columns=NAME:.metadata.name,NODE:.spec.nodeName,IMAGE:.spec.containers[0].image

# Sort
k get pods --sort-by=.metadata.creationTimestamp
k get nodes --sort-by=.status.capacity.cpu
k get pods --sort-by='.status.containerStatuses[0].restartCount'

# Output to file (exam)
k get nodes -o json > /opt/output.json
k get nodes -o jsonpath='{.items[*].status.nodeInfo.osImage}' > /opt/os.txt

# Count
k get pods -A --no-headers | wc -l

# All images running in cluster
k get pods -A -o jsonpath='{range .items[*]}{.spec.containers[*].image}{"\\n"}{end}' | sort | uniq

# Find pod's node
k get pod <name> -o jsonpath='{.spec.nodeName}'

# Decode secret
k get secret <name> -o jsonpath='{.data.password}' | base64 --decode""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 31. HELM
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('31. Helm', level=1)
doc.add_paragraph('Package manager for Kubernetes. Charts = templated YAML. Helm 3 (no Tiller). Release = deployed chart instance.')

add_code("""# Add repo
helm repo add bitnami https://charts.bitnami.com/bitnami
helm repo update

# Install
helm install my-nginx bitnami/nginx
helm install my-nginx bitnami/nginx -f values.yaml
helm install my-nginx bitnami/nginx --set service.port=8080

# List releases
helm list
helm list -A

# Upgrade
helm upgrade my-nginx bitnami/nginx --set replicaCount=3

# Rollback
helm rollback my-nginx 1

# Uninstall
helm uninstall my-nginx

# Template (render without installing)
helm template my-nginx bitnami/nginx -f values.yaml

# Show default values
helm show values bitnami/nginx""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 32. KUSTOMIZE
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('32. Kustomize', level=1)
doc.add_paragraph('Built into kubectl. Patch-based customization without templating. Uses kustomization.yaml.')

add_yaml("""# Base: kustomization.yaml
apiVersion: kustomize.config.k8s.io/v1beta1
kind: Kustomization
resources:
  - deployment.yaml
  - service.yaml

# Overlay: overlays/prod/kustomization.yaml
apiVersion: kustomize.config.k8s.io/v1beta1
kind: Kustomization
resources:
  - ../../base
patches:
  - target:
      kind: Deployment
      name: myapp
    patch: |-
      - op: replace
        path: /spec/replicas
        value: 5
images:
  - name: myapp
    newTag: v2.0""")

add_code("""kubectl apply -k overlays/prod
kubectl kustomize overlays/prod   # preview without applying""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 33. KUBEADM INSTALLATION
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('33. kubeadm Installation', level=1)

doc.add_heading('Prerequisites (All Nodes)', level=2)
add_code("""# Disable swap
sudo swapoff -a
sed -i '/ swap / s/^/#/' /etc/fstab

# Kernel modules
cat <<EOF | sudo tee /etc/modules-load.d/k8s.conf
overlay
br_netfilter
EOF
sudo modprobe overlay
sudo modprobe br_netfilter

# Sysctl
cat <<EOF | sudo tee /etc/sysctl.d/k8s.conf
net.bridge.bridge-nf-call-iptables  = 1
net.bridge.bridge-nf-call-ip6tables = 1
net.ipv4.ip_forward                 = 1
EOF
sudo sysctl --system

# Install packages
sudo apt-get update
sudo apt-get install -y kubelet kubeadm kubectl
sudo apt-mark hold kubelet kubeadm kubectl""")

doc.add_heading('Control Plane Init', level=2)
add_code("""sudo kubeadm init --pod-network-cidr=10.244.0.0/16 --apiserver-advertise-address=<master-ip>

# Copy kubeconfig
mkdir -p $HOME/.kube
sudo cp -i /etc/kubernetes/admin.conf $HOME/.kube/config
sudo chown $(id -u):$(id -g) $HOME/.kube/config

# Install CNI (e.g. Flannel)
kubectl apply -f https://raw.githubusercontent.com/flannel-io/flannel/master/Documentation/kube-flannel.yml""")

doc.add_heading('Join Workers', level=2)
add_code("""# Use token from kubeadm init output:
sudo kubeadm join <master-ip>:6443 --token <token> --discovery-token-ca-cert-hash sha256:<hash>

# If token expired:
kubeadm token create --print-join-command""")

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 34. CRDs & OPERATORS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('34. CRDs & Operators', level=1)
doc.add_paragraph('Custom Resource Definitions (CRDs) extend the API with custom types.')
doc.add_paragraph('Operators = CRD + Controller that manages lifecycle (install, upgrade, backup).')

add_code("""k get crd
k get <custom-resource-name>
k describe crd <name>""")

doc.add_heading('Extension Interfaces', level=2)
doc.add_paragraph('CNI (Container Network Interface): pod networking plugins (Calico, Flannel, Weave)')
doc.add_paragraph('CSI (Container Storage Interface): storage drivers for PVs')
doc.add_paragraph('CRI (Container Runtime Interface): container runtimes (containerd, CRI-O)')

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 35. PRACTICE SCENARIOS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('35. Practice Scenarios & Mock Questions', level=1)

scenarios = [
    ("Scenario 1: Deployment + Service",
     "Create deployment nginx-deploy --image=nginx:1.21 --replicas=3 -n app. Expose as ClusterIP service nginx-svc on port 80.",
     """k create deployment nginx-deploy --image=nginx:1.21 --replicas=3 -n app
k expose deployment nginx-deploy --name=nginx-svc --port=80 -n app
k get deploy,svc -n app"""),

    ("Scenario 2: RBAC — read-only pods",
     "Create Role allowing get,list on pods in dev namespace. Bind to user jane.",
     """k create role pod-reader --verb=get,list --resource=pods -n dev
k create rolebinding jane-pod-reader --role=pod-reader --user=jane -n dev
k auth can-i get pods --as=jane -n dev"""),

    ("Scenario 3: Static Pod",
     "Create static pod static-busybox (busybox, sleep 3600) on control plane.",
     """cat <<EOF > /etc/kubernetes/manifests/static-busybox.yaml
apiVersion: v1
kind: Pod
metadata:
  name: static-busybox
spec:
  containers:
  - name: busybox
    image: busybox:1.28
    command: ["sleep", "3600"]
EOF"""),

    ("Scenario 4: etcd Backup",
     "Snapshot etcd to /tmp/etcd-snapshot.db with correct certs.",
     """ETCDCTL_API=3 etcdctl snapshot save /tmp/etcd-snapshot.db \\
  --endpoints=https://127.0.0.1:2379 \\
  --cacert=/etc/kubernetes/pki/etcd/ca.crt \\
  --cert=/etc/kubernetes/pki/etcd/server.crt \\
  --key=/etc/kubernetes/pki/etcd/server.key"""),

    ("Scenario 5: PV + PVC + Pod",
     "PV pv-log 100Mi RWX hostPath /pv/log. PVC claim-log-1 50Mi RWX. Pod logger (nginx) mount at /var/log/nginx.",
     """# PV
apiVersion: v1
kind: PersistentVolume
metadata: { name: pv-log }
spec:
  capacity: { storage: 100Mi }
  accessModes: [ReadWriteMany]
  hostPath: { path: /pv/log }
---
# PVC
apiVersion: v1
kind: PersistentVolumeClaim
metadata: { name: claim-log-1 }
spec:
  accessModes: [ReadWriteMany]
  resources: { requests: { storage: 50Mi } }
---
# Pod
apiVersion: v1
kind: Pod
metadata: { name: logger }
spec:
  containers:
  - name: nginx
    image: nginx
    volumeMounts: [{ name: log-vol, mountPath: /var/log/nginx }]
  volumes:
  - name: log-vol
    persistentVolumeClaim: { claimName: claim-log-1 }"""),

    ("Scenario 6: NetworkPolicy — deny all + allow specific",
     "Namespace secure. Select pods app=db. Allow ingress only from app=api on TCP 5432.",
     """apiVersion: networking.k8s.io/v1
kind: NetworkPolicy
metadata:
  name: db-policy
  namespace: secure
spec:
  podSelector: { matchLabels: { app: db } }
  policyTypes: [Ingress]
  ingress:
  - from:
    - podSelector: { matchLabels: { app: api } }
    ports: [{ protocol: TCP, port: 5432 }]"""),

    ("Scenario 7: Cluster Upgrade (v1.30 → v1.31)",
     "Upgrade control plane to v1.31.0.",
     """sudo apt-mark unhold kubeadm
sudo apt-get update && sudo apt-get install -y kubeadm=1.31.0-*
sudo apt-mark hold kubeadm
sudo kubeadm upgrade plan
sudo kubeadm upgrade apply v1.31.0
kubectl drain <cp-node> --ignore-daemonsets
sudo apt-mark unhold kubelet kubectl
sudo apt-get install -y kubelet=1.31.0-* kubectl=1.31.0-*
sudo apt-mark hold kubelet kubectl
sudo systemctl daemon-reload && sudo systemctl restart kubelet
kubectl uncordon <cp-node>"""),

    ("Scenario 8: Ingress with TLS",
     "Create TLS secret webapp-tls. Ingress webapp-ingress TLS on webapp.example.com → webapp-svc:80.",
     """k create secret tls webapp-tls --cert=tls.crt --key=tls.key -n apps
# Ingress:
spec:
  tls:
  - hosts: [webapp.example.com]
    secretName: webapp-tls
  rules:
  - host: webapp.example.com
    http:
      paths:
      - path: /
        pathType: Prefix
        backend:
          service: { name: webapp-svc, port: { number: 80 } }"""),

    ("Scenario 9: Multi-container Pod Logs",
     "Stream previous logs from sidecar container in myapp-pod.",
     """kubectl logs myapp-pod -c sidecar --previous -f"""),

    ("Scenario 10: JSONPath Queries",
     "List nodes sorted by CPU. Custom columns for pods. Get InternalIP of all nodes.",
     """# Nodes sorted by CPU
k get nodes --sort-by=.status.capacity.cpu

# Custom columns
k get pods -A -o=custom-columns='NAME:.metadata.name,IMAGE:.spec.containers[*].image'

# InternalIP
k get nodes -o jsonpath='{.items[*].status.addresses[?(@.type=="InternalIP")].address}'"""),
]

for title, task, solution in scenarios:
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    run = p.add_run('Task: ')
    run.bold = True
    p.add_run(task)
    doc.add_heading('Solution:', level=3)
    add_code(solution)

add_section_break()

# ═══════════════════════════════════════════════════════════════════
# 36. EXAM TIPS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('36. Exam Tips & Common Mistakes', level=1)

doc.add_heading('Time Management', level=2)
doc.add_paragraph('2 hours, ~17 questions. Average ~7 min/question. Flag hard ones, return later. High-weight first (troubleshooting = 30%).')

doc.add_heading('Essential Bookmarks for Exam', level=2)
add_bullet('kubernetes.io/docs/reference/kubectl/cheatsheet/')
add_bullet('kubernetes.io/docs/concepts/ (Workloads, Services, Storage, Config)')
add_bullet('kubernetes.io/docs/tasks/ (Administer Cluster, Manage TLS, Configure Pods)')
add_bullet('kubernetes.io/docs/reference/ (API Reference)')

doc.add_heading('Common Mistakes to Avoid', level=2)
add_bullet('Forgetting to switch context: k config use-context <context> — each question may use different cluster')
add_bullet('Wrong namespace — always check and use -n <ns>')
add_bullet('YAML indentation errors — use k apply -f and read the error')
add_bullet('Not verifying — always k get / k describe to confirm')
add_bullet('Spending too long on one question — flag and move on')
add_bullet('Forgetting --dry-run=client -o yaml — fastest way to generate templates')

doc.add_heading('Useful One-Liners', level=2)
add_code("""# All pods with node info
k get pods -A -o wide

# Events sorted by time
k get events -A --sort-by='.lastTimestamp'

# Watch pods live
k get pods -w

# Delete stuck Terminating pod
k delete pod <name> --force --grace-period=0

# All images in cluster
k get pods -A -o jsonpath='{range .items[*]}{.spec.containers[*].image}{"\\n"}{end}' | sort | uniq

# Quick test pod
k run tmp --image=busybox:1.28 --rm -it --restart=Never -- sh

# Decode secret
k get secret <name> -o jsonpath='{.data.password}' | base64 --decode

# Find static pod path
grep staticPodPath /var/lib/kubelet/config.yaml

# Component health
k get --raw='/readyz?verbose'""")

doc.add_heading('Priority Classes (Pod Scheduling)', level=2)
add_yaml("""apiVersion: scheduling.k8s.io/v1
kind: PriorityClass
metadata: { name: high-priority }
value: 1000000
globalDefault: false
preemptionPolicy: PreemptLowerPriority

# Use in Pod:
spec:
  priorityClassName: high-priority""")

doc.add_heading('Admission Controllers', level=2)
doc.add_paragraph('Intercept API requests after auth. Validating: accept/reject (PodSecurity, ResourceQuota). Mutating: modify (DefaultStorageClass, ServiceAccount). Configure via --enable-admission-plugins on kube-apiserver.')

doc.add_heading('API Groups Quick Reference', level=2)
add_code("""# Core group (no prefix):  apiVersion: v1
# Pod, Service, ConfigMap, Secret, Namespace, PV, PVC, Endpoints

# apps group:  apiVersion: apps/v1
# Deployment, ReplicaSet, DaemonSet, StatefulSet

# batch group:  apiVersion: batch/v1
# Job, CronJob

# networking group:  apiVersion: networking.k8s.io/v1
# NetworkPolicy, Ingress

# rbac group:  apiVersion: rbac.authorization.k8s.io/v1
# Role, ClusterRole, RoleBinding, ClusterRoleBinding

# storage group:  apiVersion: storage.k8s.io/v1
# StorageClass

# policy group:  apiVersion: policy/v1
# PodDisruptionBudget

# autoscaling group:  apiVersion: autoscaling/v2
# HorizontalPodAutoscaler

# certificates group:  apiVersion: certificates.k8s.io/v1
# CertificateSigningRequest

# scheduling group:  apiVersion: scheduling.k8s.io/v1
# PriorityClass

k api-resources          # full list
k api-versions           # all API versions""")


# ═══════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'CKA_Exam_Comprehensive_CheatSheet.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
print('Done!')
